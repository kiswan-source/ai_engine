"""File writers: DOCX, PDF, HTML, TXT, JSON, XLSX, PPTX."""
import os, json
from typing import Any, Dict

from agent.tools.readers import _with_transient_retry  # Fase 15 — shared narrow-retry helper

OUTPUT_DIR = os.path.expanduser("~/ai_engine/reports")
os.makedirs(OUTPUT_DIR, exist_ok=True)

def _path(filename: str) -> str:
    return os.path.join(OUTPUT_DIR, filename) if not os.path.dirname(filename) else filename

def write_txt(filename: str, content: str) -> Dict[str, Any]:
    path = _path(filename)
    def _do():
        with open(path, "w", encoding="utf-8") as f: f.write(content)
        return {"success": True, "file": path, "filename": filename, "size": os.path.getsize(path), "type": "txt"}
    try:
        return _with_transient_retry(_do)
    except Exception as e:
        return {"success": False, "error": str(e)}

def write_json(filename: str, data: Any) -> Dict[str, Any]:
    path = _path(filename)
    def _do():
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False, default=str)
        return {"success": True, "file": path, "filename": filename, "size": os.path.getsize(path), "type": "json"}
    try:
        return _with_transient_retry(_do)
    except Exception as e:
        return {"success": False, "error": str(e)}

def write_html(filename: str, content: str) -> Dict[str, Any]:
    path = _path(filename)
    if "<html" not in content.lower():
        content = f"""<!DOCTYPE html><html lang="id"><head><meta charset="UTF-8">
<title>AI Engine Report</title>
<style>body{{font-family:'Segoe UI',sans-serif;max-width:900px;margin:40px auto;padding:0 24px;background:#f8f9fa;color:#212529;line-height:1.7}}
h1{{color:#00c896;border-bottom:2px solid #00c896;padding-bottom:8px}}h2{{color:#495057;margin-top:28px}}
table{{width:100%;border-collapse:collapse;margin:16px 0}}th,td{{border:1px solid #dee2e6;padding:8px 12px}}
th{{background:#00c896;color:white}}tr:nth-child(even){{background:#f2f2f2}}
.card{{background:white;border-radius:8px;padding:20px;margin:16px 0;box-shadow:0 1px 4px rgba(0,0,0,.1)}}
code{{background:#e9ecef;padding:2px 6px;border-radius:4px;font-family:monospace}}
pre{{background:#1e2020;color:#e8eaea;padding:16px;border-radius:8px;overflow-x:auto}}</style>
</head><body>{content}</body></html>"""
    def _do():
        with open(path, "w", encoding="utf-8") as f: f.write(content)
        return {"success": True, "file": path, "filename": filename, "size": os.path.getsize(path), "type": "html"}
    try:
        return _with_transient_retry(_do)
    except Exception as e:
        return {"success": False, "error": str(e)}

def write_docx(filename: str, title: str, content: str) -> Dict[str, Any]:
    path = _path(filename)
    def _do():
        from docx import Document
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from core.document.markdown_render import render_docx_body
        doc = Document()
        t = doc.add_heading(title, 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in t.runs: run.font.color.rgb = RGBColor(0x00, 0xc8, 0x96)
        doc.add_paragraph()
        render_docx_body(doc, content)
        # Gate 2 finding (Fase 15, 2026-08-02): save straight to `path`, unlike
        # append_pdf_section/edit_docx_section's already-established tmp-then-
        # os.replace pattern — a transient failure mid-save (now retried up to
        # 2 more times by _with_transient_retry) could otherwise leave a
        # truncated/corrupted file at `path` if every attempt failed. Closes
        # the same class of gap those two functions already closed for
        # themselves.
        tmp_path = f"{path}.tmp"
        doc.save(tmp_path)
        os.replace(tmp_path, path)
        return {"success": True, "file": path, "filename": filename, "size": os.path.getsize(path), "type": "docx"}
    try:
        return _with_transient_retry(_do)
    except Exception as e:
        return {"success": False, "error": str(e)}

def write_xlsx(filename: str, title: str, content: str) -> Dict[str, Any]:
    """Workspace Slice 3 (Fase 12) — same {filename, title, content} shape
    as write_docx/write_pdf, reusing the same markdown_render.py parse."""
    path = _path(filename)
    def _do():
        from core.document.markdown_render import render_xlsx_workbook
        wb = render_xlsx_workbook(content, default_title=title)
        tmp_path = f"{path}.tmp"  # Gate 2 finding (Fase 15) — see write_docx
        wb.save(tmp_path)
        os.replace(tmp_path, path)
        return {"success": True, "file": path, "filename": filename, "size": os.path.getsize(path), "type": "xlsx"}
    try:
        return _with_transient_retry(_do)
    except Exception as e:
        return {"success": False, "error": str(e)}

def write_pptx(filename: str, title: str, content: str) -> Dict[str, Any]:
    """Workspace Slice 3 (Fase 12) — same {filename, title, content} shape
    as write_docx/write_pdf, reusing the same markdown_render.py parse."""
    path = _path(filename)
    def _do():
        from core.document.markdown_render import render_pptx_slides
        prs = render_pptx_slides(content, default_title=title)
        tmp_path = f"{path}.tmp"  # Gate 2 finding (Fase 15) — see write_docx
        prs.save(tmp_path)
        os.replace(tmp_path, path)
        return {"success": True, "file": path, "filename": filename, "size": os.path.getsize(path), "type": "pptx"}
    try:
        return _with_transient_retry(_do)
    except Exception as e:
        return {"success": False, "error": str(e)}

def _pdf_content_styles(styles, accent) -> Dict[str, Any]:
    """The ParagraphStyle set render_pdf_story requires (h1/h2/body/bullet/
    quote/code/table_header/table_cell) — factored out of write_pdf so
    append_pdf_section (Workspace Slice 4, Fase 12) can build visually
    consistent appended pages without duplicating this construction."""
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.styles import ParagraphStyle
    return {
        "h1": ParagraphStyle("H1", parent=styles["Heading1"], fontSize=14, spaceBefore=12, spaceAfter=4),
        "h2": ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12, spaceBefore=8, spaceAfter=3),
        "body": ParagraphStyle("B", parent=styles["Normal"], fontSize=10, leading=16, alignment=TA_JUSTIFY),
        "bullet": ParagraphStyle("BU", parent=styles["Normal"], fontSize=10, leading=14, leftIndent=16),
        "quote": ParagraphStyle("Q", parent=styles["Normal"], fontSize=10, leading=14, leftIndent=16,
                                 textColor=HexColor("#495057")),
        "code": ParagraphStyle("C", parent=styles["Normal"], fontName="Courier", fontSize=9, leading=12,
                                backColor=HexColor("#f1f3f5"), borderPadding=6),
        "table_header": ParagraphStyle("TH", parent=styles["Normal"], fontSize=9, leading=12),
        "table_cell": ParagraphStyle("TC", parent=styles["Normal"], fontSize=9, leading=12),
    }

def write_pdf(filename: str, title: str, content: str) -> Dict[str, Any]:
    path = _path(filename)
    def _do():
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_CENTER
        from core.document.markdown_render import esc_pdf_markup, render_pdf_story
        tmp_path = f"{path}.tmp"  # Gate 2 finding (Fase 15) — see write_docx
        doc = SimpleDocTemplate(tmp_path, pagesize=A4, rightMargin=2.2*cm, leftMargin=2.2*cm, topMargin=2.5*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        accent = HexColor("#00c896")
        content_styles = _pdf_content_styles(styles, accent)
        s_title = ParagraphStyle("T", parent=styles["Title"], textColor=accent, fontSize=18, spaceAfter=6, alignment=TA_CENTER)
        story = [Paragraph(esc_pdf_markup(title), s_title), HRFlowable(width="100%", thickness=1.5, color=accent, spaceAfter=8), Spacer(1, 0.3*cm)]
        story.extend(render_pdf_story(content, content_styles))
        doc.build(story)
        os.replace(tmp_path, path)
        return {"success": True, "file": path, "filename": filename, "size": os.path.getsize(path), "type": "pdf"}
    try:
        return _with_transient_retry(_do)
    except Exception as e:
        return {"success": False, "error": str(e)}

def append_pdf_section(filename: str, content: str, title: str = None) -> Dict[str, Any]:
    """Workspace Slice 4 (Fase 12) — PDF half of in-place edit. Appends new
    pages built from ``content`` (optionally under a new H1-styled
    ``title``) to the END of the PDF already at ``filename``, leaving every
    existing page byte-for-byte untouched — never mid-document text
    editing. Deliberately dependency-free (pypdf + reportlab, both already
    present) rather than using PyMuPDF (AGPL-3.0-licensed, DCF-classified
    HUMAN-ONLY/HALT over the license exposure) — Owner's Gate 1 decision
    was this narrower scope instead. See
    core/document/markdown_render.py::edit_docx_section's docstring for the
    DOCX half's different, less-constrained mechanism."""
    import io

    from pypdf import PdfReader, PdfWriter
    from reportlab.lib.colors import HexColor
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

    from core.document.markdown_render import esc_pdf_markup, render_pdf_story

    path = _path(filename)
    if not os.path.exists(path):
        return {"success": False, "error": f"{filename!r} tidak ditemukan — append butuh PDF yang sudah ada."}

    def _do():
        existing_reader = PdfReader(path)
        pages_before = len(existing_reader.pages)

        buf = io.BytesIO()
        new_doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=2.2*cm, leftMargin=2.2*cm, topMargin=2.5*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        accent = HexColor("#00c896")
        story = []
        if title:
            s_title = ParagraphStyle("T", parent=styles["Heading1"], textColor=accent, fontSize=16, spaceAfter=6)
            story += [Paragraph(esc_pdf_markup(title), s_title), HRFlowable(width="100%", thickness=1, color=accent, spaceAfter=8), Spacer(1, 0.2*cm)]
        story.extend(render_pdf_story(content, _pdf_content_styles(styles, accent)))
        new_doc.build(story)
        buf.seek(0)
        new_reader = PdfReader(buf)

        writer = PdfWriter()
        for page in existing_reader.pages:
            writer.add_page(page)
        for page in new_reader.pages:
            writer.add_page(page)

        # Write-then-replace rather than writing straight into `path` — a
        # failure mid-write must never leave the original file corrupted or
        # truncated (the pre-write WorkspaceFileVersion snapshot already
        # covers "the caller changed their mind", this covers "the process
        # crashed mid-write"). Also safe to retry whole (Fase 15): nothing
        # touches `path` itself until the final os.replace.
        tmp_path = f"{path}.tmp"
        with open(tmp_path, "wb") as f:
            writer.write(f)
        os.replace(tmp_path, path)

        return {"success": True, "file": path, "filename": filename, "type": "pdf",
                "pages_before": pages_before, "pages_after": pages_before + len(new_reader.pages),
                "pages_added": len(new_reader.pages), "size": os.path.getsize(path)}
    try:
        return _with_transient_retry(_do)
    except Exception as e:
        return {"success": False, "error": str(e)}
