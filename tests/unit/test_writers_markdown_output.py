"""write_docx/write_pdf end-to-end with realistic mixed-markdown content
(Fase 11 fix) — the reported bug required calling the actual public tool
functions, not just the shared parser, since the bug could also live in how
writers.py wires the title/heading-0 block around the parsed body.
"""
from agent.tools.writers import write_docx, write_pdf

CONTENT = """### **Ringkasan Eksekutif**

Dokumen ini **penting** dan berisi *catatan*.

- Poin satu
- Poin dua

1. Langkah pertama
2. Langkah kedua

| Nama | Luas (Ha) |
| --- | --- |
| Blok A | 11.35 |
"""


def test_write_docx_produces_real_structure_no_literal_markdown(tmp_path):
    from docx import Document

    out = tmp_path / "laporan.docx"
    result = write_docx(str(out), "Laporan Uji", CONTENT)

    assert result["success"] is True
    doc = Document(str(out))
    heading = next(p for p in doc.paragraphs if p.style.name == "Heading 3")
    assert heading.text == "Ringkasan Eksekutif"
    assert heading.runs[0].bold is True
    assert len(doc.tables) == 1

    all_text = "\n".join(p.text for p in doc.paragraphs)
    for leaked in ("###", "**", "- Poin", "1. Langkah"):
        assert leaked not in all_text


def test_write_pdf_produces_valid_pdf_no_literal_markdown(tmp_path):
    from pypdf import PdfReader

    out = tmp_path / "laporan.pdf"
    result = write_pdf(str(out), "Laporan Uji", CONTENT)

    assert result["success"] is True
    reader = PdfReader(str(out))
    text = "\n".join(p.extract_text() for p in reader.pages)
    assert "Ringkasan Eksekutif" in text
    assert "Blok A" in text
    for leaked in ("###", "**", "| Nama"):
        assert leaked not in text


def test_write_pdf_title_with_markup_special_characters_does_not_break_or_drop_text(tmp_path):
    """Gate 2 fix: title went straight into ReportLab's markup parser
    unescaped — "A & B <Report>" either broke the whole write (a real
    inline tag left unclosed) or silently dropped the bracketed substring
    (parsed as an unknown empty tag), with zero error surfaced either way."""
    from pypdf import PdfReader

    out = tmp_path / "laporan.pdf"
    result = write_pdf(str(out), "A & B <Report>", "isi")

    assert result["success"] is True
    text = PdfReader(str(out)).pages[0].extract_text()
    assert "A & B <Report>" in text
