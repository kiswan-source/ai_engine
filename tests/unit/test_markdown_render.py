"""Unit tests for core/document/markdown_render.py (Fase 11 — "format
profesional" fix). agent/tools/writers.py::write_docx/write_pdf used to
hand-parse markdown line by line and left literal "###"/"**"/"|" characters
in generated documents — these tests lock in the real, structured output a
single markdown-it-py pass now produces for both formats.
"""
import pytest

from core.document.markdown_render import (
    Run,
    edit_docx_section,
    parse_markdown,
    render_docx_body,
    render_pdf_story,
    render_pptx_slides,
    render_xlsx_workbook,
)


def _pdf_styles():
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet

    base = getSampleStyleSheet()
    return {
        "h1": ParagraphStyle("H1", parent=base["Heading1"], fontSize=14, leading=18),
        "h2": ParagraphStyle("H2", parent=base["Heading2"], fontSize=12, leading=15),
        "body": ParagraphStyle("B", parent=base["Normal"], fontSize=10, leading=16),
        "bullet": ParagraphStyle("BU", parent=base["Normal"], fontSize=10, leading=14, leftIndent=16),
        "quote": ParagraphStyle("Q", parent=base["Normal"], fontSize=10, leading=14, leftIndent=16),
        "code": ParagraphStyle("C", parent=base["Normal"], fontName="Courier", fontSize=9, leading=12),
        "table_header": ParagraphStyle("TH", parent=base["Normal"], fontSize=9, leading=12),
        "table_cell": ParagraphStyle("TC", parent=base["Normal"], fontSize=9, leading=12),
    }


def test_heading_with_bold_produces_single_bold_run_not_literal_asterisks():
    """The exact bug reported live: "### **Ringkasan**" used to become a
    docx heading whose TEXT was the literal string "**Ringkasan**"."""
    blocks = parse_markdown("### **Ringkasan**")
    assert len(blocks) == 1
    assert blocks[0].kind == "heading"
    assert blocks[0].level == 3
    assert blocks[0].runs == [Run("Ringkasan", bold=True)]


def test_paragraph_with_mixed_bold_italic_code():
    blocks = parse_markdown("Ini **bold** dan *italic* serta `kode`.")
    assert len(blocks) == 1
    runs = blocks[0].runs
    texts = [(r.text, r.bold, r.italic, r.code) for r in runs]
    assert ("bold", True, False, False) in texts
    assert ("italic", False, True, False) in texts
    assert ("kode", False, False, True) in texts


def test_bullet_list_items():
    blocks = parse_markdown("- Poin satu\n- Poin dua")
    assert len(blocks) == 1
    assert blocks[0].kind == "bullet_list"
    assert [runs[0].text for _depth, runs in blocks[0].items] == ["Poin satu", "Poin dua"]
    assert [depth for depth, _runs in blocks[0].items] == [1, 1]


def test_ordered_list_items():
    blocks = parse_markdown("1. Nomor satu\n2. Nomor dua")
    assert len(blocks) == 1
    assert blocks[0].kind == "ordered_list"
    assert [runs[0].text for _depth, runs in blocks[0].items] == ["Nomor satu", "Nomor dua"]


def test_ordered_list_custom_start_number_is_preserved():
    blocks = parse_markdown("3. Tiga\n4. Empat")
    assert blocks[0].start == 3


def test_nested_bullet_list_items_are_kept_not_dropped():
    """Gate 2 regression: nested items used to vanish entirely (depth > 1
    was skipped outright), not just render unindented."""
    blocks = parse_markdown("- Item 1\n  - Nested A\n  - Nested B\n- Item 2")
    assert len(blocks) == 1
    texts = [runs[0].text for _depth, runs in blocks[0].items]
    assert texts == ["Item 1", "Nested A", "Nested B", "Item 2"]
    depths = [depth for depth, _runs in blocks[0].items]
    assert depths == [1, 2, 2, 1]


def test_image_produces_placeholder_run_not_empty():
    """Gate 2 regression: images used to produce zero runs (silently
    dropped) with no trace at all."""
    blocks = parse_markdown("![Peta lokasi](http://example.com/map.png)")
    assert len(blocks) == 1
    assert blocks[0].runs
    assert "Peta lokasi" in blocks[0].runs[0].text


def test_gfm_table_header_and_rows():
    table = "| Nama | Luas (Ha) |\n| --- | --- |\n| Blok A | 11.35 |\n| Blok B | 5.20 |"
    blocks = parse_markdown(table)
    assert len(blocks) == 1
    assert blocks[0].kind == "table"
    assert [c[0].text for c in blocks[0].header] == ["Nama", "Luas (Ha)"]
    assert [[c[0].text for c in row] for row in blocks[0].rows] == [["Blok A", "11.35"], ["Blok B", "5.20"]]


def test_blockquote():
    blocks = parse_markdown("> Catatan penting")
    assert blocks[0].kind == "blockquote"
    assert blocks[0].runs[0].text == "Catatan penting"


def test_fenced_code_block():
    blocks = parse_markdown("```\nprint('hi')\n```")
    assert blocks[0].kind == "code_block"
    assert blocks[0].text == "print('hi')"


def test_horizontal_rule():
    blocks = parse_markdown("Sebelum\n\n---\n\nSesudah")
    kinds = [b.kind for b in blocks]
    assert kinds == ["paragraph", "hr", "paragraph"]


# ─── render_docx_body: real python-docx structure, not text dump ──────────

def test_render_docx_body_produces_real_heading_and_bold_run(tmp_path):
    from docx import Document

    doc = Document()
    render_docx_body(doc, "### **Ringkasan**\n\nIsi **penting**.")

    heading = next(p for p in doc.paragraphs if p.style.name == "Heading 3")
    assert heading.text == "Ringkasan"
    assert heading.runs[0].bold is True
    # The literal syntax must never appear anywhere in the document's text.
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "###" not in all_text
    assert "**" not in all_text


def test_render_docx_body_produces_real_bullet_and_number_list_styles():
    from docx import Document

    doc = Document()
    render_docx_body(doc, "- Satu\n- Dua\n\n1. Alpha\n2. Beta")

    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    assert styles.count("List Bullet") == 2
    assert styles.count("List Number") == 2
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "- Satu" not in all_text
    assert "1. Alpha" not in all_text


def test_render_docx_body_produces_real_table():
    from docx import Document

    doc = Document()
    render_docx_body(doc, "| A | B |\n| --- | --- |\n| x | y |")

    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert [c.text for c in table.rows[0].cells] == ["A", "B"]
    assert [c.text for c in table.rows[1].cells] == ["x", "y"]
    # No paragraph anywhere should contain the raw pipe syntax.
    assert not any("|" in p.text for p in doc.paragraphs)


def test_render_docx_body_nested_list_uses_distinct_style():
    from docx import Document

    doc = Document()
    render_docx_body(doc, "- Item 1\n  - Nested A\n- Item 2")

    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    assert styles == ["List Bullet", "List Bullet 2", "List Bullet"]


# ─── render_pdf_story: heading fidelity, nested/start-numbered lists ──────

def test_render_pdf_story_differentiates_h2_h3_h4_font_sizes():
    """Gate 2 regression: every level >= 2 used to collapse into one h2
    style, so a document with H2/H3/H4 structure rendered visually flat."""
    story = render_pdf_story("## H2\n\n### H3\n\n#### H4", _pdf_styles())
    sizes = [p.style.fontSize for p in story if hasattr(p, "style")]
    assert sizes[0] > sizes[1] > sizes[2]


def test_render_pdf_story_ordered_list_respects_custom_start():
    story = render_pdf_story("3. Tiga\n4. Empat", _pdf_styles())
    texts = [p.text for p in story if hasattr(p, "text")]
    assert texts[0].startswith("3.")
    assert texts[1].startswith("4.")


def test_render_pdf_story_nested_list_item_is_kept_and_indented():
    story = render_pdf_story("- Item 1\n  - Nested A\n- Item 2", _pdf_styles())
    paragraphs = [p for p in story if hasattr(p, "text")]
    texts = [p.text for p in paragraphs]
    assert any("Nested A" in t for t in texts)
    nested = next(p for p in paragraphs if "Nested A" in p.text)
    top = next(p for p in paragraphs if "Item 1" in p.text)
    assert nested.style.leftIndent > top.style.leftIndent


# ─── render_xlsx_workbook / render_pptx_slides (Workspace Slice 3, Fase 12) ─

def test_render_xlsx_workbook_h1_starts_new_sheet():
    md = "# Ringkasan\nIsi ringkasan.\n\n# Data\nIsi data."
    wb = render_xlsx_workbook(md, default_title="Laporan")
    assert wb.sheetnames == ["Ringkasan", "Data"]


def test_render_xlsx_workbook_content_before_first_heading_uses_default_title():
    md = "Paragraf pembuka tanpa heading.\n\n# Detail\nIsi."
    wb = render_xlsx_workbook(md, default_title="Laporan")
    assert wb.sheetnames == ["Laporan", "Detail"]
    assert wb["Laporan"]["A1"].value == "Paragraf pembuka tanpa heading."


def test_render_xlsx_workbook_no_heading_at_all_uses_one_default_sheet():
    wb = render_xlsx_workbook("Cuma satu paragraf.", default_title="Laporan")
    assert wb.sheetnames == ["Laporan"]


def test_render_xlsx_workbook_table_becomes_real_rows():
    md = "# Data\n| Nama | Luas |\n| --- | --- |\n| Blok A | 11.35 |\n| Blok B | 5.20 |"
    wb = render_xlsx_workbook(md, default_title="Laporan")
    ws = wb["Data"]
    rows = [[c.value for c in row] for row in ws.iter_rows()]
    assert rows == [["Nama", "Luas"], ["Blok A", "11.35"], ["Blok B", "5.20"]]
    header_cell = ws["A1"]
    assert header_cell.font.bold is True


def test_render_xlsx_workbook_duplicate_heading_names_get_unique_sheet_names():
    md = "# Data\nSatu.\n\n# Data\nDua."
    wb = render_xlsx_workbook(md, default_title="Laporan")
    assert wb.sheetnames == ["Data", "Data-2"]


def test_render_xlsx_workbook_heavy_duplication_stays_within_31_char_limit():
    """Adversarial review finding: a fixed base[:28] truncation only stays
    within Excel's 31-char sheet-name limit while the dedup counter is 1-2
    digits — 100+ duplicate headings used to silently produce an invalid,
    spec-violating (32+ char) sheet name."""
    md = "\n\n".join(f"# {'X' * 31}" for _ in range(105))
    wb = render_xlsx_workbook(md, default_title="Laporan")
    assert all(len(name) <= 31 for name in wb.sheetnames)
    assert len(wb.sheetnames) == len(set(wb.sheetnames))  # still unique


def test_render_xlsx_workbook_consecutive_headings_with_no_body_are_not_dropped():
    """Adversarial review finding: an H1 immediately followed by another H1
    (no body in between — a realistic shape of model output, e.g. an
    outline heading not yet filled in) used to vanish with zero trace,
    since the old logic only kept non-empty groups."""
    md = "# Ringkasan\nIsi ringkasan lengkap.\n\n# Rekomendasi\n\n# Penutup\nTerima kasih."
    wb = render_xlsx_workbook(md, default_title="Laporan")
    assert wb.sheetnames == ["Ringkasan", "Rekomendasi", "Penutup"]


def test_render_pptx_slides_consecutive_headings_with_no_body_are_not_dropped():
    md = "# Ringkasan\nIsi ringkasan lengkap.\n\n# Rekomendasi\n\n# Penutup\nTerima kasih."
    prs = render_pptx_slides(md, default_title="Laporan")
    titles = [s.shapes.title.text for s in prs.slides]
    assert titles == ["Ringkasan", "Rekomendasi", "Penutup"]


def test_render_xlsx_workbook_sanitizes_invalid_sheet_name_characters():
    md = "# A/B:C*D\nIsi."
    wb = render_xlsx_workbook(md, default_title="Laporan")
    assert wb.sheetnames == ["A-B-C-D"]


def test_render_xlsx_workbook_neutralizes_formula_injection_in_table_cells():
    # Gate 3 (AEGIS audit, 2026-07-23) — CWE-1236: a cell value starting with
    # =/+/-/@ must never reach the xlsx as a live formula. openpyxl itself
    # auto-detects a bare "=..." string as a formula cell (data_type "f");
    # the leading apostrophe forces text and is the standard mitigation.
    md = (
        "# Sheet1\n\n"
        "| Item | Val |\n| --- | --- |\n"
        "| A | =2+2 |\n| B | +CMD |\n| C | -1 |\n| D | @SUM(A1) |\n| E | normal |\n"
    )
    wb = render_xlsx_workbook(md, default_title="Sheet1")
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    values_by_item = {r[0]: r[1] for r in rows[1:]}
    assert values_by_item["A"] == "'=2+2"
    assert values_by_item["B"] == "'+CMD"
    assert values_by_item["C"] == "'-1"
    assert values_by_item["D"] == "'@SUM(A1)"
    assert values_by_item["E"] == "normal"
    # The critical property: openpyxl must never classify this as a formula
    # cell (data_type 'f') — confirmed by construction since the stored
    # value no longer starts with "=".
    for cell in ws["B"]:
        assert cell.data_type != "f"


def test_render_pptx_slides_h1_starts_new_slide():
    md = "# Ringkasan\nIsi ringkasan.\n\n# Data\nIsi data."
    prs = render_pptx_slides(md, default_title="Laporan")
    titles = [s.shapes.title.text for s in prs.slides]
    assert titles == ["Ringkasan", "Data"]


def test_render_pptx_slides_body_gets_bullet_lines():
    """Gate 2 fix: this used to only assert substring presence, which
    passed regardless of whether a "•" marker was actually there — list
    items rendered as bare text with zero marker at all, unlike the xlsx/
    PDF renderers built from the same IR. Assert the real marker now."""
    md = "# Ringkasan\n- Poin satu\n- Poin dua"
    prs = render_pptx_slides(md, default_title="Laporan")
    slide = next(iter(prs.slides))
    paragraphs = slide.placeholders[1].text_frame.paragraphs
    assert [p.text for p in paragraphs] == ["• Poin satu", "• Poin dua"]
    assert all(p.level == 0 for p in paragraphs)


def test_render_pptx_slides_ordered_list_gets_number_markers():
    md = "# Ringkasan\n1. Langkah satu\n2. Langkah dua"
    prs = render_pptx_slides(md, default_title="Laporan")
    slide = next(iter(prs.slides))
    texts = [p.text for p in slide.placeholders[1].text_frame.paragraphs]
    assert texts == ["1. Langkah satu", "2. Langkah dua"]


def test_render_pptx_slides_nested_list_uses_paragraph_level():
    md = "# Ringkasan\n- Item 1\n  - Nested A\n- Item 2"
    prs = render_pptx_slides(md, default_title="Laporan")
    slide = next(iter(prs.slides))
    paragraphs = slide.placeholders[1].text_frame.paragraphs
    levels = {p.text: p.level for p in paragraphs}
    assert levels["• Item 1"] == 0
    assert levels["• Nested A"] == 1
    assert levels["• Item 2"] == 0


def test_render_pptx_slides_table_gets_its_own_slide_with_real_table_shape():
    md = "# Data\n| A | B |\n| --- | --- |\n| 1 | 2 |"
    prs = render_pptx_slides(md, default_title="Laporan")
    slides = list(prs.slides)
    assert len(slides) == 2
    assert slides[1].shapes.title.text == "Data — Tabel"
    table_shapes = [s for s in slides[1].shapes if s.has_table]
    assert len(table_shapes) == 1
    table = table_shapes[0].table
    assert [c.text for c in table.rows[0].cells] == ["A", "B"]
    assert [c.text for c in table.rows[1].cells] == ["1", "2"]


def test_render_pptx_slides_no_heading_at_all_uses_one_default_slide():
    prs = render_pptx_slides("Cuma satu paragraf.", default_title="Laporan")
    assert len(list(prs.slides)) == 1
    assert next(iter(prs.slides)).shapes.title.text == "Laporan"


# ─── edit_docx_section (Workspace Slice 4, in-place edit, Fase 12) ────────

def _doc_with_sections():
    from docx import Document

    doc = Document()
    doc.add_heading("Judul", 0)
    render_docx_body(
        doc,
        "# Ringkasan\nIsi ringkasan lama.\n\n# Data\nIsi data lama.\n\n# Penutup\nIsi penutup.",
    )
    return doc


def test_edit_docx_section_rejects_blank_heading_text():
    """Gate 2 fix: the chat-facing wrapper already guards against a blank
    heading, but this function is itself public — a blank/whitespace-only
    heading_text used to match nothing, fall to the "not found" branch, and
    silently append a real Heading 1 paragraph with EMPTY text."""
    doc = _doc_with_sections()
    with pytest.raises(ValueError):
        edit_docx_section(doc, "   ", "isi baru")


def test_edit_docx_section_replaces_only_the_matched_section_body():
    doc = _doc_with_sections()
    action = edit_docx_section(doc, "Data", "Isi data BARU.")
    assert action == "edited"
    texts = [(p.style.name, p.text) for p in doc.paragraphs]
    assert ("Normal", "Isi ringkasan lama.") in texts  # untouched
    assert ("Normal", "Isi penutup.") in texts  # untouched
    assert ("Normal", "Isi data lama.") not in texts  # replaced
    assert ("Normal", "Isi data BARU.") in texts
    # Heading order/structure preserved.
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert headings == ["Ringkasan", "Data", "Penutup"]


def test_edit_docx_section_heading_match_is_case_insensitive_and_trims_whitespace():
    doc = _doc_with_sections()
    action = edit_docx_section(doc, "  data  ", "Isi BARU.")
    assert action == "edited"
    assert any(p.text == "Isi BARU." for p in doc.paragraphs)


def test_edit_docx_section_appends_when_heading_not_found():
    doc = _doc_with_sections()
    action = edit_docx_section(doc, "Rekomendasi", "Isi rekomendasi baru.")
    assert action == "appended"
    texts = [p.text for p in doc.paragraphs]
    assert texts[-2:] == ["Rekomendasi", "Isi rekomendasi baru."]
    assert doc.paragraphs[-2].style.name == "Heading 1"


def test_edit_docx_section_removes_nested_subsections_under_the_target():
    """A deeper heading (H2 under the target H1) is part of the section
    being replaced, not a boundary — it must be removed along with
    everything else under the target heading."""
    from docx import Document

    doc = Document()
    render_docx_body(
        doc,
        "# Data\nIntro.\n\n## Sub A\nIsi sub A.\n\n## Sub B\nIsi sub B.\n\n# Penutup\nIsi penutup.",
    )
    action = edit_docx_section(doc, "Data", "Isi data baru, tanpa subsection.")
    assert action == "edited"
    texts = [p.text for p in doc.paragraphs]
    assert "Sub A" not in texts and "Sub B" not in texts
    assert "Isi sub A." not in texts and "Isi sub B." not in texts
    assert "Isi data baru, tanpa subsection." in texts
    assert "Penutup" in texts and "Isi penutup." in texts  # untouched, still last


def test_edit_docx_section_last_section_has_no_boundary_and_still_works():
    doc = _doc_with_sections()
    action = edit_docx_section(doc, "Penutup", "Penutup BARU.")
    assert action == "edited"
    texts = [p.text for p in doc.paragraphs]
    assert texts[-1] == "Penutup BARU."
    assert "Isi penutup." not in texts


def test_edit_docx_section_last_section_preserves_document_sectPr(tmp_path):
    """Adversarial review finding: editing the LAST section (no boundary
    heading found) used to delete every body child up to the end of the
    document, including the trailing <w:sectPr> (page size/margins/
    header-footer refs) — corrupting the document (doc.sections goes from
    1 to 0), surviving save+reload. Regression test checks the real
    save/reload round trip, not just in-memory state."""
    from docx import Document

    doc = _doc_with_sections()
    assert len(doc.sections) == 1
    edit_docx_section(doc, "Penutup", "Penutup BARU.")
    assert len(doc.sections) == 1

    out = tmp_path / "edited.docx"
    doc.save(str(out))
    reloaded = Document(str(out))
    assert len(reloaded.sections) == 1


def test_edit_docx_section_matches_first_occurrence_when_heading_text_repeats():
    from docx import Document

    doc = Document()
    render_docx_body(doc, "# Data\nSection pertama.\n\n# Lain\nIsi lain.\n\n# Data\nSection kedua.")
    edit_docx_section(doc, "Data", "Sudah diedit.")
    texts = [p.text for p in doc.paragraphs]
    assert "Sudah diedit." in texts
    assert "Section pertama." not in texts
    assert "Section kedua." in texts  # second "Data" section untouched


def test_edit_docx_section_ignores_non_heading_paragraph_with_matching_text():
    """A bold body paragraph that happens to say "Data" must not be
    mistaken for the "Data" heading — only a paragraph actually styled
    Heading 1-4 counts as a match."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Data")  # plain "Normal" style, not a heading
    render_docx_body(doc, "# Data\nIsi section.")
    action = edit_docx_section(doc, "Data", "Isi BARU.")
    assert action == "edited"
    texts = [(p.style.name, p.text) for p in doc.paragraphs]
    assert ("Normal", "Data") in texts  # the non-heading paragraph survives untouched
    assert ("Normal", "Isi BARU.") in texts
    assert ("Normal", "Isi section.") not in texts


def test_edit_docx_section_table_is_positioned_inside_the_replaced_section():
    from docx.text.paragraph import Paragraph

    doc = _doc_with_sections()
    edit_docx_section(doc, "Data", "Isi baru.\n\n| A | B |\n| --- | --- |\n| 1 | 2 |")

    body_children = list(doc.element.body)
    tags = [el.tag.rsplit("}", 1)[-1] for el in body_children]
    data_idx = next(
        i for i, el in enumerate(body_children)
        if tags[i] == "p" and Paragraph(el, doc).text == "Data"
    )
    penutup_idx = next(
        i for i, el in enumerate(body_children)
        if tags[i] == "p" and Paragraph(el, doc).text == "Penutup"
    )
    tbl_idx = tags.index("tbl")
    assert data_idx < tbl_idx < penutup_idx


def test_edit_docx_section_disambiguates_by_heading_level():
    from docx import Document

    doc = Document()
    render_docx_body(doc, "# Data\nIsi H1.\n\n## Data\nIsi H2.")
    edit_docx_section(doc, "Data", "Isi H2 BARU.", heading_level=2)
    texts = [(p.style.name, p.text) for p in doc.paragraphs]
    assert ("Normal", "Isi H1.") in texts  # H1 section untouched
    assert ("Normal", "Isi H2 BARU.") in texts
